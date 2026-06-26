#!/usr/bin/env python3
"""Core utilities for grabbing YouTube transcripts and making study files."""

from __future__ import annotations

import json
import os
import re
import subprocess
import sys
import time
import asyncio
import tempfile
import zipfile
from dataclasses import asdict, dataclass
from hashlib import sha1
from pathlib import Path
from typing import Any
from urllib.parse import parse_qs, urlparse

import requests

DEEPSEEK_API_URL = "https://api.deepseek.com/chat/completions"
DEFAULT_OUTPUT_DIR = "output"
SUBTITLE_LANG_PRIORITY = ["zh-Hans", "zh", "en", "en-US"]
BATCH_SIZE = 20
REQUEST_SLEEP_SECONDS = 1.0
PROJECT_ROOT = Path(__file__).resolve().parent
DEFAULT_ENV_FILE = PROJECT_ROOT / ".env"
ARTICLE_USER_AGENT = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
    "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/126.0 Safari/537.36"
)
TTS_VOICES = {
    "us-female": "en-US-JennyNeural",
    "us-male": "en-US-GuyNeural",
    "uk-female": "en-GB-SoniaNeural",
    "uk-male": "en-GB-RyanNeural",
    "default": "en-US-JennyNeural",
}
TTS_RATES = {
    "0.85": "-15%",
    "1": "+0%",
    "1.15": "+15%",
}


@dataclass
class TranscriptSnippet:
    text: str
    start: float
    duration: float
    text_zh: str = ""

    @classmethod
    def from_mapping(cls, data: dict) -> "TranscriptSnippet":
        return cls(
            text=str(data.get("text", "")),
            start=float(data.get("start", 0)),
            duration=float(data.get("duration", 0)),
            text_zh=str(data.get("text_zh", "")),
        )

    def to_dict(self) -> dict:
        data = asdict(self)
        if not data["text_zh"]:
            data.pop("text_zh")
        return data


def extract_video_id(url: str) -> str | None:
    """Extract the 11-character YouTube video id from common video URLs."""
    value = (url or "").strip()
    if re.fullmatch(r"[0-9A-Za-z_-]{11}", value):
        return value

    parsed = urlparse(value)
    host = parsed.netloc.lower().removeprefix("www.")
    path = parsed.path.strip("/")

    if host in {"youtube.com", "m.youtube.com", "music.youtube.com"}:
        query_video_id = parse_qs(parsed.query).get("v", [None])[0]
        if query_video_id and re.fullmatch(r"[0-9A-Za-z_-]{11}", query_video_id):
            return query_video_id

        parts = path.split("/")
        if len(parts) >= 2 and parts[0] in {"shorts", "embed", "live"}:
            candidate = parts[1]
            if re.fullmatch(r"[0-9A-Za-z_-]{11}", candidate):
                return candidate

    if host == "youtu.be":
        candidate = path.split("/")[0]
        if re.fullmatch(r"[0-9A-Za-z_-]{11}", candidate):
            return candidate

    match = re.search(r"(?:v=|youtu\.be/|/shorts/|/embed/|/live/)([0-9A-Za-z_-]{11})", value)
    return match.group(1) if match else None


def is_single_video(url: str) -> bool:
    return extract_video_id(url) is not None


def is_youtube_source(url: str) -> bool:
    value = (url or "").strip()
    if re.fullmatch(r"[0-9A-Za-z_-]{11}", value):
        return True

    host = urlparse(value).netloc.lower().removeprefix("www.")
    return host in {"youtube.com", "m.youtube.com", "music.youtube.com", "youtu.be"}


def get_deepseek_api_key(env_file: str | Path = DEFAULT_ENV_FILE) -> str | None:
    api_key = os.environ.get("DEEPSEEK_API_KEY", "").strip()
    if api_key:
        return api_key

    path = Path(env_file)
    if not path.exists():
        return None

    for line in path.read_text(encoding="utf-8").splitlines():
        stripped = line.strip()
        if not stripped or stripped.startswith("#"):
            continue
        key, separator, value = stripped.partition("=")
        if separator and key.strip() == "DEEPSEEK_API_KEY":
            return value.strip().strip("'\"") or None

    return None


def _run_yt_dlp(url: str) -> subprocess.CompletedProcess:
    commands = [
        [sys.executable, "-m", "yt_dlp", "--flat-playlist", "--print", "id", url],
        ["yt-dlp", "--flat-playlist", "--print", "id", url],
    ]
    last_error: Exception | None = None

    for command in commands:
        try:
            return subprocess.run(
                command,
                capture_output=True,
                text=True,
                check=True,
                timeout=120,
                encoding="utf-8",
            )
        except FileNotFoundError as exc:
            last_error = exc
        except subprocess.CalledProcessError:
            raise

    raise RuntimeError("未找到 yt-dlp，请先运行：py -m pip install -r requirements.txt") from last_error


def list_video_ids_from_channel_or_playlist(url: str) -> list[str]:
    """Return video ids from a channel, playlist, or other yt-dlp-supported URL."""
    try:
        result = _run_yt_dlp(url)
    except subprocess.CalledProcessError as exc:
        message = exc.stderr.strip() or exc.stdout.strip() or str(exc)
        raise RuntimeError(f"yt-dlp 解析失败：{message}") from exc

    video_ids: list[str] = []
    seen: set[str] = set()
    for line in result.stdout.splitlines():
        candidate = line.strip()
        if re.fullmatch(r"[0-9A-Za-z_-]{11}", candidate) and candidate not in seen:
            video_ids.append(candidate)
            seen.add(candidate)
    return video_ids


def resolve_video_ids(url: str, limit: int | None = None) -> list[str]:
    if is_single_video(url):
        video_ids = [extract_video_id(url)]
    else:
        video_ids = list_video_ids_from_channel_or_playlist(url)

    resolved = [video_id for video_id in video_ids if video_id]
    return resolved[:limit] if limit else resolved


def fetch_transcript_for_video(video_id: str) -> list[TranscriptSnippet] | None:
    """Fetch one video's transcript, preferring Chinese then English."""
    try:
        from youtube_transcript_api import YouTubeTranscriptApi
        from youtube_transcript_api._errors import NoTranscriptFound, TranscriptsDisabled, VideoUnavailable
    except ImportError as exc:
        raise RuntimeError("缺少 youtube-transcript-api，请先运行：py -m pip install -r requirements.txt") from exc

    ytt_api = YouTubeTranscriptApi()
    try:
        transcript_list = ytt_api.list(video_id)
    except TranscriptsDisabled:
        return None
    except VideoUnavailable:
        return None
    except Exception as exc:
        raise RuntimeError(f"{video_id}: 获取字幕列表失败：{exc}") from exc

    transcript = None
    for lang in SUBTITLE_LANG_PRIORITY:
        try:
            transcript = transcript_list.find_transcript([lang])
            break
        except NoTranscriptFound:
            continue

    if transcript is None:
        available = list(transcript_list)
        transcript = available[0] if available else None

    if transcript is None:
        return None

    try:
        fetched = transcript.fetch()
    except Exception as exc:
        raise RuntimeError(f"{video_id}: 抓取字幕内容失败：{exc}") from exc

    return [
        TranscriptSnippet(text=item.text, start=item.start, duration=item.duration)
        for item in fetched
    ]


def fetch_article(url: str) -> dict[str, Any]:
    """Extract readable text from a public article page."""
    response = requests.get(
        url,
        headers={"User-Agent": ARTICLE_USER_AGENT, "Accept": "text/html,application/xhtml+xml"},
        timeout=30,
    )
    response.raise_for_status()
    if not response.encoding:
        response.encoding = "utf-8"

    try:
        from bs4 import BeautifulSoup
    except ImportError as exc:
        raise RuntimeError("缺少 beautifulsoup4，无法解析文章。请先运行：py -m pip install -r requirements.txt") from exc

    soup = BeautifulSoup(response.text, "html.parser")
    for selector in "script style noscript nav footer header form aside".split():
        for node in soup.select(selector):
            node.decompose()

    title = _first_text(
        soup.select_one("article h1"),
        soup.select_one("h1"),
        soup.select_one('meta[property="og:title"]'),
        soup.select_one("title"),
    )
    author = _first_text(
        soup.select_one('meta[name="author"]'),
        soup.select_one('meta[property="article:author"]'),
        soup.select_one('[rel="author"]'),
        soup.select_one(".author"),
        soup.select_one(".byline"),
    )

    container = soup.select_one("article") or soup.select_one("main") or soup.body or soup
    paragraphs: list[str] = []
    seen: set[str] = set()
    for node in container.select("p"):
        text = " ".join(node.get_text(" ", strip=True).split())
        if len(text) < 20 or text in seen:
            continue
        paragraphs.append(text)
        seen.add(text)

    if not paragraphs:
        raise RuntimeError("没有从这个网页识别到可用正文。可能需要登录、页面禁止抓取，或它不是文章页。")

    return {
        "title": title or url,
        "author": author,
        "paragraphs": paragraphs,
        "source_url": url,
    }


def fetch_video_metadata(video_id: str) -> dict[str, str]:
    source_url = f"https://www.youtube.com/watch?v={video_id}"
    metadata = {"title": "", "author": "", "source_url": source_url}
    try:
        response = requests.get(
            "https://www.youtube.com/oembed",
            params={"url": source_url, "format": "json"},
            headers={"User-Agent": ARTICLE_USER_AGENT},
            timeout=10,
        )
        response.raise_for_status()
        data = response.json()
    except Exception:
        return metadata

    metadata["title"] = str(data.get("title") or "").strip()
    metadata["author"] = str(data.get("author_name") or "").strip()
    return metadata


def _first_text(*nodes) -> str:
    for node in nodes:
        if not node:
            continue
        value = node.get("content") if hasattr(node, "get") and node.get("content") else node.get_text(" ", strip=True)
        text = " ".join(str(value or "").split())
        if text:
            return text
    return ""


def translate_batch(texts: list[str], api_key: str) -> list[str]:
    numbered_input = "\n".join(f"{index + 1}. {text}" for index, text in enumerate(texts))
    prompt = (
        "你是专业字幕翻译。请将下面编号的英文口语字幕逐条翻译成简体中文，"
        "保持口语化、自然、适合跟读学习。不要漏号，不要合并句子，不要增删句子数量。"
        "严格按照“编号. 译文”的格式逐行输出，不要输出其他说明文字。\n\n"
        f"{numbered_input}"
    )

    response = requests.post(
        DEEPSEEK_API_URL,
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        json={
            "model": "deepseek-chat",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.3,
        },
        timeout=60,
    )
    response.raise_for_status()
    content = response.json()["choices"][0]["message"]["content"]

    translations: dict[int, str] = {}
    for line in content.splitlines():
        match = re.match(r"^(\d+)\.\s*(.*)$", line.strip())
        if match:
            translations[int(match.group(1))] = match.group(2)

    return [translations.get(index + 1, "") for index in range(len(texts))]


def _parse_json_object(content: str) -> dict:
    text = content.strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text)
        text = re.sub(r"\s*```$", "", text)

    start = text.find("{")
    end = text.rfind("}")
    if start != -1 and end != -1 and end > start:
        text = text[start : end + 1]

    return json.loads(text)


def analyze_vocab_term(term: str, context: str, api_key: str) -> dict[str, str]:
    clean_term = " ".join(str(term or "").strip().split())
    clean_context = " ".join(str(context or "").strip().split())
    if not clean_term:
        raise ValueError("请选择要标注的英文词或短语。")

    prompt = (
        "你是英语精读老师。请根据上下文解释这个英文生词或短语，输出严格 JSON，"
        "不要 Markdown，不要额外说明。字段必须是："
        "term、part_of_speech、definition_zh、example_zh。"
        "part_of_speech 用 n./v./adj./adv./phr. 等简短形式；"
        "definition_zh 用中文，适合放在生词栏，尽量简洁。\n\n"
        f"生词：{clean_term}\n"
        f"上下文：{clean_context}"
    )

    response = requests.post(
        DEEPSEEK_API_URL,
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        json={
            "model": "deepseek-chat",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.2,
        },
        timeout=60,
    )
    response.raise_for_status()
    content = response.json()["choices"][0]["message"]["content"]
    parsed = _parse_json_object(content)

    return {
        "term": str(parsed.get("term") or clean_term).strip(),
        "part_of_speech": str(parsed.get("part_of_speech") or "phr.").strip(),
        "definition_zh": str(parsed.get("definition_zh") or "").strip(),
        "example_zh": str(parsed.get("example_zh") or "").strip(),
    }


def generate_tts_audio(
    text: str,
    output_root: str = DEFAULT_OUTPUT_DIR,
    voice_profile: str = "us-female",
    rate: str = "1",
) -> str:
    clean_text = " ".join(str(text or "").split())
    if not clean_text:
        raise ValueError("朗读文本为空。")
    if len(clean_text) > 3500:
        raise ValueError("单次朗读文本过长，请分段朗读。")

    voice = TTS_VOICES.get(str(voice_profile), TTS_VOICES["default"])
    edge_rate = TTS_RATES.get(str(rate), "+0%")
    cache_key = sha1(f"{voice}|{edge_rate}|{clean_text}".encode("utf-8")).hexdigest()
    out_dir = Path(output_root) / "tts"
    out_dir.mkdir(parents=True, exist_ok=True)
    audio_path = out_dir / f"{cache_key}.mp3"
    if audio_path.exists() and audio_path.stat().st_size > 0:
        return str(audio_path)

    try:
        import edge_tts
    except ImportError as exc:
        raise RuntimeError("缺少 edge-tts，无法生成朗读音频。请先运行：py -m pip install -r requirements.txt") from exc

    async def _save_audio() -> None:
        communicate = edge_tts.Communicate(clean_text, voice, rate=edge_rate)
        await communicate.save(str(audio_path))

    asyncio.run(_save_audio())
    return str(audio_path)


def translate_transcript(snippets: list[TranscriptSnippet], api_key: str) -> list[TranscriptSnippet]:
    translated_all: list[str] = []
    texts = [snippet.text for snippet in snippets]

    for start in range(0, len(texts), BATCH_SIZE):
        batch = texts[start : start + BATCH_SIZE]
        translated_all.extend(translate_batch(batch, api_key))
        time.sleep(REQUEST_SLEEP_SECONDS)

    for snippet, text_zh in zip(snippets, translated_all):
        snippet.text_zh = text_zh
    return snippets


def format_timestamp(seconds: float) -> str:
    milliseconds = int(round(float(seconds) * 1000))
    hours, remainder = divmod(milliseconds, 3_600_000)
    minutes, remainder = divmod(remainder, 60_000)
    secs, millis = divmod(remainder, 1000)
    return f"{hours:02}:{minutes:02}:{secs:02}.{millis:03}"


def build_bilingual_text(snippets: list[TranscriptSnippet]) -> str:
    blocks: list[str] = []
    for snippet in snippets:
        block = [f"[{format_timestamp(snippet.start)}]", f"EN: {snippet.text}"]
        if snippet.text_zh:
            block.append(f"ZH: {snippet.text_zh}")
        blocks.append("\n".join(block))
    return "\n\n".join(blocks) + ("\n" if blocks else "")


def save_outputs(video_id: str, snippets: list[TranscriptSnippet] | list[dict], output_root: str = DEFAULT_OUTPUT_DIR) -> dict[str, str]:
    normalized = [item if isinstance(item, TranscriptSnippet) else TranscriptSnippet.from_mapping(item) for item in snippets]
    out_dir = Path(output_root) / video_id
    out_dir.mkdir(parents=True, exist_ok=True)

    json_path = out_dir / "transcript.json"
    txt_path = out_dir / "transcript.txt"
    zh_path = out_dir / "transcript_zh.txt"
    practice_path = out_dir / "practice_bilingual.txt"

    json_path.write_text(
        json.dumps([snippet.to_dict() for snippet in normalized], ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    txt_path.write_text("\n".join(snippet.text for snippet in normalized) + "\n", encoding="utf-8")
    practice_path.write_text(build_bilingual_text(normalized), encoding="utf-8")

    written = {
        "json": str(json_path),
        "txt": str(txt_path),
        "practice": str(practice_path),
    }

    if any(snippet.text_zh for snippet in normalized):
        zh_path.write_text("\n".join(snippet.text_zh for snippet in normalized) + "\n", encoding="utf-8")
        written["zh"] = str(zh_path)

    return written


def save_study_sheet_html(html: str, output_root: str = DEFAULT_OUTPUT_DIR, export_format: str = "html") -> str:
    content = str(html or "").strip()
    if not content:
        raise ValueError("精读稿内容为空，无法导出。")

    out_dir = Path(output_root)
    out_dir.mkdir(parents=True, exist_ok=True)
    normalized_format = str(export_format or "html").lower().strip()

    if normalized_format == "html":
        path = out_dir / "study-sheet.html"
        path.write_text(content, encoding="utf-8")
    elif normalized_format == "docx":
        path = out_dir / "study-sheet.docx"
        _write_study_sheet_docx(content, path)
    elif normalized_format == "pdf":
        path = out_dir / "study-sheet.pdf"
        _write_study_sheet_pdf(content, path)
    else:
        raise ValueError("导出格式只支持 html、docx 或 pdf。")

    return str(path)


def _extract_study_sheet_model(html: str) -> dict[str, Any]:
    try:
        from bs4 import BeautifulSoup
    except ImportError as exc:
        raise RuntimeError("缺少 beautifulsoup4，无法导出精读稿。") from exc

    soup = BeautifulSoup(html, "html.parser")
    sheet = soup.select_one(".study-sheet") or soup
    title = _node_text(sheet.select_one(".sheet-titlebar h3")) or _node_text(sheet.select_one("h3")) or "双语精读稿"
    brand = _node_text(sheet.select_one(".sheet-brand")) or "Study Sheet"
    meta = _extract_sheet_meta(sheet)
    if meta.get("题目"):
        title = meta["题目"]
    video = _node_text(sheet.select_one(".sheet-titlebar p"))

    paragraphs = []
    rows = sheet.select(".sheet-row")
    line_sources = rows or sheet.select(".sheet-line")
    for source in line_sources:
        line = source.select_one(".sheet-line") if rows else source
        english = _node_text(line.select_one(".sheet-en"))
        english_runs = _extract_marked_runs(line.select_one(".sheet-en"))
        chinese = _node_text(line.select_one(".sheet-zh"))
        if english or chinese:
            row_vocab = []
            for card in source.select(".vocab-card"):
                term = _node_text(card.select_one(".vocab-term"))
                definition = _node_text(card.select_one("p"))
                example = _node_text(card.select_one("small"))
                if term or definition:
                    row_vocab.append({"term": term, "definition": definition, "example": example})
            paragraphs.append({"english": english, "english_runs": english_runs, "chinese": chinese, "vocab": row_vocab})

    vocab = []
    for card in sheet.select(".vocab-card"):
        term = _node_text(card.select_one(".vocab-term"))
        definition = _node_text(card.select_one("p"))
        example = _node_text(card.select_one("small"))
        if term or definition:
            vocab.append({"term": term, "definition": definition, "example": example})

    if not vocab:
        vocab = [entry for paragraph in paragraphs for entry in paragraph.get("vocab", [])]

    return {"brand": brand, "title": title, "video": video, "meta": meta, "paragraphs": paragraphs, "vocab": vocab}


def _extract_sheet_meta(sheet) -> dict[str, str]:
    meta: dict[str, str] = {}
    for row in sheet.select(".sheet-meta-row"):
        label = _node_text(row.select_one("dt")).rstrip("：:")
        value = _node_text(row.select_one("dd"))
        if label and value:
            meta[label] = value
    return meta


def _extract_marked_runs(node) -> list[dict[str, Any]]:
    if not node:
        return []

    try:
        from bs4 import NavigableString
    except ImportError:
        NavigableString = str

    runs: list[dict[str, Any]] = []

    def walk(current, marked: bool = False) -> None:
        if isinstance(current, NavigableString):
            text = str(current)
            if text:
                if runs and runs[-1]["marked"] == marked:
                    runs[-1]["text"] += text
                else:
                    runs.append({"text": text, "marked": marked})
            return

        classes = current.get("class", []) if hasattr(current, "get") else []
        next_marked = marked or "vocab-mark" in classes
        for child in getattr(current, "children", []):
            walk(child, next_marked)

    walk(node)
    normalized: list[dict[str, Any]] = []
    for run in runs:
        text = re.sub(r"\s+", " ", run["text"])
        if text:
            normalized.append({"text": text, "marked": bool(run["marked"])})
    return normalized


def _node_text(node) -> str:
    return " ".join(node.get_text(" ", strip=True).split()) if node else ""


def _shade_docx_run(run, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    r_pr = run._r.get_or_add_rPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), fill)
    r_pr.append(shade)


def _set_docx_table_borders_none(table) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl_pr = table._tbl.tblPr
    existing = tbl_pr.first_child_found_in("w:tblBorders")
    if existing is not None:
        tbl_pr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)
    tbl_pr.append(borders)


def _set_docx_cell_border(cell, edge: str, color: str = "111827", size: str = "12") -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    element = borders.find(qn(f"w:{edge}"))
    if element is None:
        element = OxmlElement(f"w:{edge}")
        borders.append(element)
    element.set(qn("w:val"), "single")
    element.set(qn("w:sz"), size)
    element.set(qn("w:space"), "0")
    element.set(qn("w:color"), color)


def _write_study_sheet_docx(html: str, path: Path) -> None:
    try:
        from docx import Document
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:
        raise RuntimeError("缺少 python-docx，无法导出 DOCX。") from exc

    data = _extract_study_sheet_model(html)
    document = Document()
    section = document.sections[0]
    section.left_margin = Pt(54)
    section.right_margin = Pt(54)
    section.top_margin = Pt(54)
    section.bottom_margin = Pt(54)

    brand = document.add_paragraph()
    brand_run = brand.add_run(data["brand"])
    brand_run.bold = True
    brand_run.font.size = Pt(30)
    brand_run.font.color.rgb = RGBColor(220, 38, 38)

    title = document.add_heading(data["title"], level=1)
    for label in ("题目", "作者", "网址"):
        value = data.get("meta", {}).get(label, "")
        if value:
            row = document.add_paragraph()
            key = row.add_run(f"{label}：")
            key.bold = True
            row.add_run(value)
    if data["video"]:
        document.add_paragraph(data["video"])

    for paragraph in data["paragraphs"]:
        table = document.add_table(rows=1, cols=2)
        table.autofit = False
        _set_docx_table_borders_none(table)
        left_cell, right_cell = table.rows[0].cells
        left_cell.width = Inches(4.65)
        right_cell.width = Inches(2.05)
        left_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        right_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        _set_docx_cell_border(left_cell, "right")

        en = left_cell.paragraphs[0]
        english_runs = paragraph.get("english_runs") or [{"text": paragraph["english"], "marked": False}]
        for run_data in english_runs:
            run = en.add_run(run_data["text"])
            run.bold = True
            run.font.size = Pt(15)
            if run_data.get("marked"):
                _shade_docx_run(run, "FED7AA")
        if paragraph["chinese"]:
            zh = left_cell.add_paragraph(paragraph["chinese"])
            for run in zh.runs:
                run.font.size = Pt(12)
            zh.paragraph_format.space_after = Pt(10)
        for entry in paragraph.get("vocab", []):
            note = right_cell.add_paragraph()
            term = note.add_run(entry["term"])
            term.bold = True
            term.font.size = Pt(13)
            _shade_docx_run(term, "E6F4D7")
            note.add_run(f"  {entry['definition']}")
            if entry["example"]:
                example = right_cell.add_paragraph(entry["example"])
                for run in example.runs:
                    run.font.size = Pt(9)
        document.add_paragraph()

    if data["vocab"]:
        document.add_page_break()
        document.add_heading("生词合集", level=1)
        for entry in data["vocab"]:
            p = document.add_paragraph()
            term = p.add_run(entry["term"])
            term.bold = True
            p.add_run(f"  {entry['definition']}")
            if entry["example"]:
                document.add_paragraph(entry["example"])

    document.save(path)


def _write_study_sheet_pdf(html: str, path: Path) -> None:
    try:
        import fitz
    except ImportError as exc:
        raise RuntimeError("缺少 PyMuPDF，无法导出 PDF。") from exc

    data = _extract_study_sheet_model(html)
    font_file = _find_cjk_font()
    document = fitz.open()
    page = document.new_page(width=595, height=842)
    margin = 48
    y = margin
    line_height = 16

    def ensure_space(height: float) -> None:
        nonlocal page, y
        if y + height > 800:
            page = document.new_page(width=595, height=842)
            y = margin

    def write_line(text: str, size: int = 11, color=(0, 0, 0), bold: bool = False) -> None:
        nonlocal y
        for line in _wrap_export_text(text, 58 if size <= 12 else 36):
            ensure_space(line_height + 4)
            page.insert_textbox(
                fitz.Rect(margin, y, 547, y + line_height + 4),
                line,
                fontsize=size,
                fontfile=font_file,
                fontname="cjk" if font_file else "helv",
                color=color,
            )
            y += line_height + (2 if not bold else 4)

    def text_width(text: str, size: int) -> float:
        units = sum(2 if ord(char) > 127 else 1 for char in str(text))
        return units * size * 0.32

    def write_span_line(spans: list[dict[str, Any]], x: float, top: float, size: int = 13) -> None:
        cursor = x
        baseline = top + size + 3
        for span in spans:
            text = span.get("text", "")
            if not text:
                continue
            width = text_width(text, size)
            if span.get("marked"):
                page.draw_rect(
                    fitz.Rect(cursor - 1, top + 1, cursor + width + 1, top + size + 5),
                    color=None,
                    fill=(1.0, 0.84, 0.67),
                )
            page.insert_text(
                fitz.Point(cursor, baseline),
                text,
                fontsize=size,
                fontfile=font_file,
                fontname="cjk" if font_file else "helv",
                color=(0, 0, 0),
            )
            cursor += width

    def write_text_at(lines: list[str], x: float, top: float, size: int = 11, color=(0, 0, 0)) -> float:
        cursor_y = top
        for line in lines:
            page.insert_textbox(
                fitz.Rect(x, cursor_y, x + 330, cursor_y + line_height + 6),
                line,
                fontsize=size,
                fontfile=font_file,
                fontname="cjk" if font_file else "helv",
                color=color,
            )
            cursor_y += line_height + 4
        return cursor_y

    write_line(data["brand"], size=26, color=(0.86, 0.15, 0.15), bold=True)
    write_line(data["title"], size=18, bold=True)
    for label in ("题目", "作者", "网址"):
        value = data.get("meta", {}).get(label, "")
        if value:
            write_line(f"{label}：{value}", size=10, color=(0.28, 0.33, 0.42))
    if data["video"]:
        write_line(data["video"], size=10, color=(0.28, 0.33, 0.42))
    y += 12

    for paragraph in data["paragraphs"]:
        if paragraph["english"]:
            write_line(paragraph["english"], size=13, bold=True)
        if paragraph["chinese"]:
            write_line(paragraph["chinese"], size=11)
        for entry in paragraph.get("vocab", []):
            write_line(f"旁注：{entry['term']}  {entry['definition']}", size=9, color=(0.25, 0.32, 0.45))
            if entry["example"]:
                write_line(entry["example"], size=8, color=(0.39, 0.45, 0.55))
        y += 8

    if data["vocab"]:
        ensure_space(80)
        y += 12
        write_line("生词合集", size=17, bold=True)
        for entry in data["vocab"]:
            write_line(f"{entry['term']}  {entry['definition']}", size=11)
            if entry["example"]:
                write_line(entry["example"], size=9, color=(0.39, 0.45, 0.55))
            y += 4

    document.save(path)
    document.close()


def _write_study_sheet_docx(html: str, path: Path) -> None:
    try:
        from docx import Document
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:
        raise RuntimeError("缺少 python-docx，无法导出 DOCX。") from exc

    data = _extract_study_sheet_model(html)
    document = Document()
    section = document.sections[0]
    section.left_margin = Pt(54)
    section.right_margin = Pt(54)
    section.top_margin = Pt(54)
    section.bottom_margin = Pt(54)

    brand = document.add_paragraph()
    brand_run = brand.add_run(data["brand"])
    brand_run.bold = True
    brand_run.font.size = Pt(30)
    brand_run.font.color.rgb = RGBColor(220, 38, 38)

    document.add_heading(data["title"], level=1)
    for label in ("题目", "作者", "网址"):
        value = data.get("meta", {}).get(label, "")
        if value:
            row = document.add_paragraph()
            key = row.add_run(f"{label}：")
            key.bold = True
            row.add_run(value)
    if data["video"]:
        document.add_paragraph(data["video"])

    for paragraph in data["paragraphs"]:
        table = document.add_table(rows=1, cols=2)
        table.autofit = False
        _set_docx_table_borders_none(table)
        left_cell, right_cell = table.rows[0].cells
        left_cell.width = Inches(4.65)
        right_cell.width = Inches(2.05)
        left_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        right_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        _set_docx_cell_border(left_cell, "right")

        en = left_cell.paragraphs[0]
        english_runs = paragraph.get("english_runs") or [{"text": paragraph.get("english", ""), "marked": False}]
        for run_data in english_runs:
            run = en.add_run(run_data["text"])
            run.bold = True
            run.font.size = Pt(15)
            if run_data.get("marked"):
                _shade_docx_run(run, "FED7AA")
        if paragraph.get("chinese"):
            zh = left_cell.add_paragraph(paragraph["chinese"])
            for run in zh.runs:
                run.font.size = Pt(12)
            zh.paragraph_format.space_after = Pt(10)

        for entry in paragraph.get("vocab", []):
            note = right_cell.add_paragraph()
            term = note.add_run(entry["term"])
            term.bold = True
            term.font.size = Pt(13)
            _shade_docx_run(term, "E6F4D7")
            note.add_run(f"  {entry['definition']}")
            if entry.get("example"):
                example = right_cell.add_paragraph(entry["example"])
                for run in example.runs:
                    run.font.size = Pt(9)
        document.add_paragraph()

    document.save(path)


def _write_study_sheet_pdf(html: str, path: Path) -> None:
    try:
        import fitz
    except ImportError as exc:
        raise RuntimeError("缺少 PyMuPDF，无法导出 PDF。") from exc

    data = _extract_study_sheet_model(html)
    font_file = _find_cjk_font()
    document = fitz.open()
    page = document.new_page(width=595, height=842)
    margin = 48
    y = margin
    line_height = 16

    def ensure_space(height: float) -> None:
        nonlocal page, y
        if y + height > 800:
            page = document.new_page(width=595, height=842)
            y = margin

    def text_width(text: str, size: int) -> float:
        units = sum(2 if ord(char) > 127 else 1 for char in str(text))
        return units * size * 0.32

    def draw_text(text: str, x: float, top: float, size: int = 11, color=(0, 0, 0), width: float = 490) -> float:
        lines = _wrap_export_text(text, max(12, int(width / (size * 0.32))))
        cursor_y = top
        for line in lines:
            page.insert_textbox(
                fitz.Rect(x, cursor_y, x + width, cursor_y + line_height + 6),
                line,
                fontsize=size,
                fontfile=font_file,
                fontname="cjk" if font_file else "helv",
                color=color,
            )
            cursor_y += line_height + 4
        return cursor_y

    def draw_run_line(spans: list[dict[str, Any]], x: float, top: float, size: int = 13) -> None:
        cursor = x
        baseline = top + size + 3
        for span in spans:
            text = span.get("text", "")
            if not text:
                continue
            width = text_width(text, size)
            if span.get("marked"):
                page.draw_rect(
                    fitz.Rect(cursor - 1, top + 1, cursor + width + 1, top + size + 5),
                    color=None,
                    fill=(1.0, 0.84, 0.67),
                )
            page.insert_text(
                fitz.Point(cursor, baseline),
                text,
                fontsize=size,
                fontfile=font_file,
                fontname="cjk" if font_file else "helv",
                color=(0, 0, 0),
            )
            cursor += width

    y = draw_text(data["brand"], margin, y, size=26, color=(0.86, 0.15, 0.15))
    y = draw_text(data["title"], margin, y + 4, size=18, color=(0, 0, 0))
    for label in ("题目", "作者", "网址"):
        value = data.get("meta", {}).get(label, "")
        if value:
            y = draw_text(f"{label}：{value}", margin, y, size=10, color=(0.28, 0.33, 0.42))
    if data["video"]:
        y = draw_text(data["video"], margin, y, size=10, color=(0.28, 0.33, 0.42))
    y += 12

    for paragraph in data["paragraphs"]:
        left_x = margin
        divider_x = 392
        right_x = 408
        english_lines = _wrap_export_runs(
            paragraph.get("english_runs") or [{"text": paragraph.get("english", ""), "marked": False}],
            44,
        )
        chinese_lines = _wrap_export_text(paragraph.get("chinese", ""), 34) if paragraph.get("chinese") else []
        vocab_height = 0
        for entry in paragraph.get("vocab", []):
            vocab_height += 24 + len(_wrap_export_text(entry["definition"], 20)) * 14
            if entry.get("example"):
                vocab_height += len(_wrap_export_text(entry["example"], 20)) * 12
        left_height = len(english_lines) * 20 + len(chinese_lines) * 18 + 12
        row_height = max(left_height, vocab_height, 34)
        ensure_space(row_height + 18)
        row_top = y
        page.draw_line(fitz.Point(divider_x, row_top), fitz.Point(divider_x, row_top + row_height), color=(0.07, 0.09, 0.12), width=0.8)

        cursor_y = row_top
        for spans in english_lines:
            draw_run_line(spans, left_x, cursor_y, size=13)
            cursor_y += 20
        if chinese_lines:
            for line in chinese_lines:
                page.insert_textbox(
                    fitz.Rect(left_x, cursor_y, divider_x - 12, cursor_y + line_height + 6),
                    line,
                    fontsize=11,
                    fontfile=font_file,
                    fontname="cjk" if font_file else "helv",
                    color=(0, 0, 0),
                )
                cursor_y += 18

        side_y = row_top
        for entry in paragraph.get("vocab", []):
            term = entry["term"]
            term_width = text_width(term, 11)
            page.draw_rect(
                fitz.Rect(right_x, side_y + 1, right_x + term_width + 8, side_y + 17),
                color=None,
                fill=(0.90, 0.96, 0.84),
            )
            page.insert_text(
                fitz.Point(right_x + 3, side_y + 14),
                term,
                fontsize=11,
                fontfile=font_file,
                fontname="cjk" if font_file else "helv",
                color=(0, 0, 0),
            )
            side_y += 21
            side_y = draw_text(entry["definition"], right_x, side_y, size=8, width=135)
            if entry.get("example"):
                side_y = draw_text(entry["example"], right_x, side_y, size=7, color=(0.39, 0.45, 0.55), width=135)
            side_y += 6

        y = row_top + row_height + 14

    document.save(path)
    document.close()


def _find_cjk_font() -> str | None:
    candidates = [
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simsun.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/arialuni.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return str(candidate)
    return None


def _wrap_export_runs(runs: list[dict[str, Any]], max_units: int) -> list[list[dict[str, Any]]]:
    lines: list[list[dict[str, Any]]] = []
    current: list[dict[str, Any]] = []
    units = 0

    for run in runs:
        marked = bool(run.get("marked"))
        tokens = re.findall(r"\S+\s*", str(run.get("text", "")))
        for token in tokens:
            token_units = sum(2 if ord(char) > 127 else 1 for char in token)
            if current and units + token_units > max_units:
                lines.append(current)
                current = []
                units = 0
            current.append({"text": token, "marked": marked})
            units += token_units

    if current:
        lines.append(current)
    return lines or [[]]


def _wrap_export_text(text: str, max_units: int) -> list[str]:
    lines: list[str] = []
    current = ""
    units = 0
    for char in str(text):
        char_units = 2 if ord(char) > 127 else 1
        if current and units + char_units > max_units:
            lines.append(current.rstrip())
            current = char
            units = char_units
        else:
            current += char
            units += char_units
    if current.strip():
        lines.append(current.rstrip())
    return lines or [""]


def _find_browser_executable() -> str | None:
    candidates = [
        Path("C:/Program Files/Google/Chrome/Application/chrome.exe"),
        Path("C:/Program Files (x86)/Google/Chrome/Application/chrome.exe"),
        Path("C:/Program Files/Microsoft/Edge/Application/msedge.exe"),
        Path("C:/Program Files (x86)/Microsoft/Edge/Application/msedge.exe"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return str(candidate)
    return None


def _write_pdf_from_html_print(html: str, path: Path) -> None:
    browser = _find_browser_executable()
    if not browser:
        raise RuntimeError("未找到 Chrome 或 Edge，无法按 HTML 原样打印 PDF。")

    with tempfile.TemporaryDirectory() as tmp:
        tmp_dir = Path(tmp)
        html_path = tmp_dir / "study-sheet.html"
        profile_dir = tmp_dir / "profile"
        html_path.write_text(html, encoding="utf-8")
        command = [
            browser,
            "--headless=new",
            "--disable-gpu",
            "--no-sandbox",
            f"--user-data-dir={profile_dir}",
            "--print-to-pdf-no-header",
            f"--print-to-pdf={path}",
            html_path.as_uri(),
        ]
        subprocess.run(command, check=True, capture_output=True, text=True, timeout=120)

    if not path.exists() or path.stat().st_size == 0:
        raise RuntimeError("浏览器 PDF 导出失败。")


def _write_docx_from_html_package(html: str, path: Path) -> None:
    content_types = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Default Extension="html" ContentType="text/html"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>"""
    rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>"""
    document_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <w:body>
    <w:altChunk r:id="htmlChunk"/>
    <w:sectPr>
      <w:pgSz w:w="12240" w:h="15840"/>
      <w:pgMar w:top="720" w:right="720" w:bottom="720" w:left="720" w:header="360" w:footer="360" w:gutter="0"/>
    </w:sectPr>
  </w:body>
</w:document>"""
    document_rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="htmlChunk" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/aFChunk" Target="afchunk.html"/>
</Relationships>"""

    path.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", content_types)
        archive.writestr("_rels/.rels", rels)
        archive.writestr("word/document.xml", document_xml)
        archive.writestr("word/_rels/document.xml.rels", document_rels)
        archive.writestr("word/afchunk.html", html)


# Final export implementations: prefer the real HTML rendering over hand-built approximations.
def _write_study_sheet_docx(html: str, path: Path) -> None:
    _write_docx_from_html_package(html, path)


def _write_study_sheet_pdf(html: str, path: Path) -> None:
    _write_pdf_from_html_print(html, path)


def load_cached_bilingual_output(video_id: str, output_root: str = DEFAULT_OUTPUT_DIR) -> list[TranscriptSnippet] | None:
    json_path = Path(output_root) / video_id / "transcript.json"
    if not json_path.exists():
        return None

    try:
        snippets = [TranscriptSnippet.from_mapping(item) for item in json.loads(json_path.read_text(encoding="utf-8"))]
    except (OSError, json.JSONDecodeError, TypeError, ValueError):
        return None

    return snippets if snippets and any(snippet.text_zh for snippet in snippets) else None


def build_process_result(
    video_id: str,
    snippets: list[TranscriptSnippet],
    files: dict[str, str],
    source_type: str = "video",
    title: str = "",
    author: str = "",
    source_url: str = "",
    media_embed_url: str | None = None,
) -> dict:
    if media_embed_url is None and source_type == "video":
        media_embed_url = f"https://www.youtube.com/embed/{video_id}"

    return {
        "video_id": video_id,
        "ok": True,
        "source_type": source_type,
        "title": title,
        "author": author,
        "source_url": source_url,
        "media_embed_url": media_embed_url or "",
        "count": len(snippets),
        "files": files,
        "preview": [snippet.to_dict() for snippet in snippets[:80]],
        "full": [snippet.to_dict() for snippet in snippets],
    }


def process_video(
    video_id: str,
    output_root: str,
    translate: bool = False,
    api_key: str | None = None,
    metadata: dict[str, str] | None = None,
) -> dict:
    metadata = metadata or {"title": "", "author": "", "source_url": f"https://www.youtube.com/watch?v={video_id}"}
    if translate:
        cached = load_cached_bilingual_output(video_id, output_root)
        if cached:
            files = save_outputs(video_id, cached, output_root)
            return build_process_result(
                video_id,
                cached,
                files,
                title=metadata.get("title", ""),
                author=metadata.get("author", ""),
                source_url=metadata.get("source_url", ""),
            )

    snippets = fetch_transcript_for_video(video_id)
    if not snippets:
        return {"video_id": video_id, "ok": False, "error": "没有可用字幕，或字幕被作者关闭。"}

    if translate:
        if not api_key:
            return {"video_id": video_id, "ok": False, "error": "未设置 DEEPSEEK_API_KEY，无法翻译。"}
        snippets = translate_transcript(snippets, api_key)

    files = save_outputs(video_id, snippets, output_root)
    return build_process_result(
        video_id,
        snippets,
        files,
        title=metadata.get("title", ""),
        author=metadata.get("author", ""),
        source_url=metadata.get("source_url", ""),
    )


def process_article(url: str, output_root: str, translate: bool = False, api_key: str | None = None) -> dict:
    article = fetch_article(url)
    article_id = "article_" + sha1(url.encode("utf-8")).hexdigest()[:12]
    snippets = [
        TranscriptSnippet(text=text, start=float(index), duration=0)
        for index, text in enumerate(article["paragraphs"])
    ]

    if translate:
        if not api_key:
            return {"video_id": article_id, "ok": False, "error": "未设置 DEEPSEEK_API_KEY，无法翻译。"}

        translations: list[str] = []
        for start in range(0, len(snippets), BATCH_SIZE):
            batch = [snippet.text for snippet in snippets[start : start + BATCH_SIZE]]
            translations.extend(translate_batch(batch, api_key))
            time.sleep(REQUEST_SLEEP_SECONDS)

        for snippet, text_zh in zip(snippets, translations):
            snippet.text_zh = text_zh

    files = save_outputs(article_id, snippets, output_root)
    return build_process_result(
        article_id,
        snippets,
        files,
        source_type="article",
        title=str(article.get("title") or ""),
        author=str(article.get("author") or ""),
        source_url=str(article.get("source_url") or url),
        media_embed_url="",
    )


def process_url(url: str, output_root: str = DEFAULT_OUTPUT_DIR, translate: bool = False, limit: int | None = None) -> dict:
    api_key = get_deepseek_api_key()
    warnings: list[str] = []
    if translate and not api_key:
        translate = False
        warnings.append("未设置 DEEPSEEK_API_KEY，已自动跳过中文翻译。")

    if not is_youtube_source(url):
        results = [process_article(url, output_root, translate=translate, api_key=api_key)]
        return {
            "ok": any(result.get("ok") for result in results),
            "video_ids": [],
            "results": results,
            "warnings": warnings,
            "output_root": str(Path(output_root)),
        }

    video_ids = resolve_video_ids(url, limit=limit)
    results = [
        process_video(
            video_id,
            output_root,
            translate=translate,
            api_key=api_key,
            metadata={**fetch_video_metadata(video_id), "source_url": url if len(video_ids) == 1 else f"https://www.youtube.com/watch?v={video_id}"},
        )
        for video_id in video_ids
    ]
    return {
        "ok": any(result.get("ok") for result in results),
        "video_ids": video_ids,
        "results": results,
        "warnings": warnings,
        "output_root": str(Path(output_root)),
    }
